from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Backend Development Report: Naryn Tourism Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Lead info
    p = doc.add_paragraph()
    run = p.add_run('Role: Lead Backend Developer\nProject: Naryn Tourism Multi-Media Platform\nArchitecture: RESTful API with Node.js & Express')
    run.bold = True
    
    # Sections
    sections = [
        ("1. Executive Summary", "The backend architecture for the Naryn Tourism platform was designed to provide a secure, scalable, and high-performance foundation for a multi-media tourism marketplace. The system facilitates role-based interactions between tourists, hosts, and administrators, supporting complex workflows such as multi-media listing creation, content modification requests, and a robust booking engine."),
        
        ("2. Technology Stack", "The platform utilizes a modern JavaScript-based stack for seamless integration and development velocity:\n"
                               "- Runtime Environment: Node.js\n"
                               "- Web Framework: Express.js\n"
                               "- Database: SQLite3 (Serverless, file-based relational database)\n"
                               "- Security: JWT for stateless session management, Bcrypt.js for password encryption.\n"
                               "- Communication: CORS enabled for frontend-backend requests.\n"
                               "- Environment: Dotenv for configuration."),
        
        ("3. Database Architecture & Schema Design", "A relational schema was implemented to ensure data integrity. Key tables include:\n"
                                                     "- Users: Profiles and roles (Tourist, Host, Admin).\n"
                                                     "- Listings: Property data with JSON array support for images/videos.\n"
                                                     "- Content Requests: Managed host updates requiring admin review.\n"
                                                     "- Bookings: Transaction records between tourists and listings."),
        
        ("4. Key Functional Implementations", "4.1 Multi-Media Handling: Configured to accept up to 200MB payloads for Base64 media.\n"
                                             "4.2 RBAC: Custom middleware (auth, checkRole) for strict authorization.\n"
                                             "4.3 Content Lifecycle: 'Pending -> Approved/Rejected' workflow with automatic data merging on approval."),
        
        ("5. API Endpoints Overview", "Core endpoints include /api/auth/register, /api/listings (POST/GET), /api/requests (POST), /api/bookings (POST), and admin status updates."),
        
        ("6. Quality Assurance & Testing", "A comprehensive automated testing strategy was implemented using test-api.js, covering unit and integration tests across all core features."),
        
        ("7. Conclusion", "The backend system successfully delivers a robust and secure API. By prioritizing multi-media flexibility and strict role-based security, the architecture is well-positioned for future scaling.")
    ]
    
    for heading, content in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(content)
        
    doc.save('Backend_Development_Report.docx')
    print('Report generated successfully: Backend_Development_Report.docx')

if __name__ == "__main__":
    create_report()
