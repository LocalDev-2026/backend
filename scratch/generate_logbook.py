from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_logbook():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Project Logbook & Final Development Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('Naryn Tourism Platform: Digitalizing Nomad Hospitality', 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n' * 2)
    
    # 1. Project Introduction
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "The Naryn Tourism Platform is a specialized full-stack marketplace designed to bridge the gap between "
        "international travelers and local hospitality providers in the Naryn region of Kyrgyzstan. The platform "
        "allows hosts to create detailed listings with high-resolution images and videos, while providing tourists "
        "with a seamless booking experience."
    )
    
    doc.add_heading('1.1 Why Naryn? Importance for Locals', level=2)
    doc.add_paragraph(
        "Naryn is a region of immense natural beauty and authentic nomadic culture, yet many local guesthouse "
        "owners and tour guides remain digitally invisible to the global market. This project is vital for Naryn locals because:\n"
        "• Economic Empowerment: It allows locals to bypass high-commission third-party agencies and connect directly with tourists.\n"
        "• Digital Visibility: High-quality multi-media support (5 images, 3 videos per listing) allows hosts to showcase the unique aesthetic of yurt stays and mountain tours.\n"
        "• Cultural Preservation: By facilitating authentic tourism, the platform helps sustain traditional livelihoods.\n"
        "• Language Accessibility: Supporting Kyrgyz, English, Russian, and Turkish ensures that local hosts can manage their businesses in their native tongue while reaching a global audience."
    )
    
    # 2. Team Roles & Responsibilities
    doc.add_heading('2. Team Roles & Contributions', level=1)
    
    # Project Manager
    doc.add_heading('2.1 Project Manager', level=2)
    doc.add_paragraph(
        "Responsibilities:\n"
        "• Requirements Engineering: Defined the core features, including the multi-media upload limits and multilingual support.\n"
        "• Timeline Management: Coordinated the transition from development to final QA audit.\n"
        "• Quality Assurance: Directed the end-to-end testing of the registration, listing, and approval workflows.\n"
        "• Final Integration: Ensured the frontend and backend repositories were synchronized and pushed to GitHub for deployment."
    )
    
    # Backend Developer
    doc.add_heading('2.2 Backend Developer', level=2)
    doc.add_paragraph(
        "Responsibilities:\n"
        "• API Architecture: Built a robust RESTful API using Node.js and Express.\n"
        "• Database Design: Designed the relational SQLite schema for users, listings, bookings, and content requests.\n"
        "• Security: Implemented JWT-based authentication and Role-Based Access Control (RBAC).\n"
        "• Media Infrastructure: Configured the system to handle 200MB payloads for Base64 image and video data storage."
    )
    
    # Frontend Developer
    doc.add_heading('2.3 Frontend Developer', level=2)
    doc.add_paragraph(
        "Responsibilities:\n"
        "• UI/UX Design: Developed a responsive React interface using modern CSS and Lucide icons.\n"
        "• State Management: Implemented complex form handling for multi-file media uploads with real-time previews.\n"
        "• Internationalization: Integrated i18next to support 4 languages (EN, KY, RU, TR) across all components.\n"
        "• Media Gallery: Created dynamic rendering components for listing images and HTML5 video players."
    )
    
    # 3. Development Log (Logbook)
    doc.add_heading('3. Project Development Log (Logbook)', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase'
    hdr_cells[1].text = 'Activities'
    hdr_cells[2].text = 'Status'
    
    log_data = [
        ("Phase 1: Planning", "Market research in Naryn, defining UI/UX requirements, and database schema design.", "Completed"),
        ("Phase 2: Core Dev", "Backend API setup, JWT auth implementation, and basic frontend routing.", "Completed"),
        ("Phase 3: Multi-Media", "Integration of 5-image/3-video upload support and Base64 conversion logic.", "Completed"),
        ("Phase 4: Multilingual", "Translation of the entire platform into KY, EN, RU, TR and i18n integration.", "Completed"),
        ("Phase 5: Admin Logic", "Creation of the Admin Dashboard and Content Request approval workflow.", "Completed"),
        ("Phase 6: QA & Audit", "Full-stack testing, bug fixing, and final synchronization with GitHub.", "Completed")
    ]
    
    for phase, activity, status in log_data:
        row_cells = table.add_row().cells
        row_cells[0].text = phase
        row_cells[1].text = activity
        row_cells[2].text = status
        
    # 4. Conclusion
    doc.add_heading('4. Conclusion', level=1)
    doc.add_paragraph(
        "The Naryn Tourism Platform is now a fully functional, stable, and localized solution. By empowering "
        "local hosts with digital tools, we are fostering a more inclusive and prosperous tourism ecosystem in Naryn."
    )
    
    doc.save('Naryn_Tourism_Project_Logbook.docx')
    print('Logbook generated: Naryn_Tourism_Project_Logbook.docx')

if __name__ == "__main__":
    create_logbook()
